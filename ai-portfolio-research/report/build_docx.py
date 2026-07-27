"""Chuyển BAO_CAO_HOAN_CHINH.md -> BAO_CAO_HOAN_CHINH.docx với style học thuật chuyên nghiệp.

Chạy:  python report/build_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "BAO_CAO_HOAN_CHINH.md"
OUT_PATH = ROOT / "BAO_CAO_HOAN_CHINH.docx"
FIG_DIR = ROOT / "figures"

ACCENT = RGBColor(0x1E, 0x3A, 0x5F)   # xanh navy chuyên nghiệp
ACCENT2 = RGBColor(0x2E, 0x5E, 0x8C)
GREY = RGBColor(0x55, 0x55, 0x55)

AUTHOR = "Quách Thành Long"
MSSV = "88241020109"
COURSE = "Phân tích đầu tư nâng cao"
DATE = "27/07/2026"


# ---------------------------------------------------------------- helpers --

def set_base_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(8)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after, keep in [
        ("Heading 1", 18, ACCENT, 26, 12, True),
        ("Heading 2", 14, ACCENT2, 16, 8, True),
        ("Heading 3", 12.5, ACCENT2, 10, 6, True),
    ]:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = keep
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)


def add_page_number_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{AUTHOR} · {MSSV} · {COURSE}   |   Trang ")
    run.font.size = Pt(9)
    run.font.color.rgb = GREY

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run2 = p.add_run()
    run2._r.append(fld_begin)
    run2._r.append(instr)
    run2._r.append(fld_sep)
    run2._r.append(fld_end)
    run2.font.size = Pt(9)
    run2.font.color.rgb = GREY


def add_toc_field(doc: Document):
    """Chèn trường Mục lục tự động (Word: nhấn F9 hoặc 'Update Field' để hiện đầy đủ)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "(Click chuột phải → Update Field để hiển thị Mục lục đầy đủ)"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(txt)
    run._r.append(fld_end)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def add_inline_runs(paragraph, text: str):
    """Diễn giải **bold**, `code`, *italic* trong 1 dòng text -> runs có style tương ứng."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0xA0, 0x30, 0x30)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            paragraph.add_run(part)


def strip_md_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)


def add_para(doc: Document, text: str, style=None, align=None, italic=False, size=None, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    add_inline_runs(p, strip_md_links(text))
    for r in p.runs:
        if italic:
            r.italic = True
        if bold:
            r.bold = True
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(header))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        add_inline_runs(p, strip_md_links(h))
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1E3A5F")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            add_inline_runs(p, strip_md_links(val))
            for r in p.runs:
                r.font.size = Pt(10.5)
    doc.add_paragraph()
    return table


# ------------------------------------------------------------------ cover --

def build_cover(doc: Document):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ỨNG DỤNG TRÍ TUỆ NHÂN TẠO\nTRONG QUẢN LÝ DANH MỤC ĐẦU TƯ")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = ACCENT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "Từ lý thuyết danh mục hiện đại đến Deep Learning: thiết kế, huấn luyện\n"
        "và đánh giá một hệ thống predict-then-optimize ngoài mẫu"
    )
    r2.italic = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = ACCENT2

    for _ in range(4):
        doc.add_paragraph()

    info = [
        ("Môn học", COURSE),
        ("Sinh viên thực hiện", AUTHOR),
        ("MSSV", MSSV),
        ("Ngày hoàn thành", DATE),
    ]
    for label, val in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}:  ")
        r.bold = True
        r.font.size = Pt(13)
        r2 = p.add_run(val)
        r2.font.size = Pt(13)

    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Mã nguồn, dữ liệu và toàn bộ kết quả thực nghiệm đi kèm báo cáo được lưu trữ tại\n"
        "thư mục ai-portfolio-research/ — có thể chạy lại độc lập, tái lập 100% số liệu trình bày."
    )
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = GREY

    doc.add_page_break()


# ------------------------------------------------------------- md parsing --

def parse_table_block(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    cells = [
        [c.strip() for c in r.strip("|").split("|")]
        for r in rows
        if not re.match(r"^\|?\s*:?-{2,}", r.split("|")[1] if "|" in r else "")
    ]
    return cells, i


def is_block_start(line: str) -> bool:
    s = line.strip()
    return (
        not s
        or s.startswith("```")
        or s.startswith("#")
        or s.startswith("|")
        or s.startswith("![")
        or s.startswith("> ")
        or s.startswith("---")
        or s.startswith("<div")
        or s.startswith("</div")
        or s.startswith("&nbsp;")
        or re.match(r"^\d+\.\s", s)
        or s.startswith("- ")
    )


def consume_para_lines(lines: list[str], start: int) -> tuple[str, int]:
    """Gộp dòng `start` với các dòng tiếp theo cùng thuộc 1 đoạn/mục (soft-wrap) -> (text, next_index)."""
    buf = [lines[start].strip()]
    j = start + 1
    while j < len(lines) and lines[j].strip() and not is_block_start(lines[j]):
        buf.append(lines[j].strip())
        j += 1
    return " ".join(buf), j


def render_body(doc: Document, md_text: str):
    lines = md_text.split("\n")
    i = 0
    in_code = False
    code_buf: list[str] = []
    skip_manual_toc = False

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Gộp các dòng liên tiếp (không rỗng, không phải block đặc biệt) thành 1 đoạn văn
        if not in_code and line.strip() and not is_block_start(line):
            text, j = consume_para_lines(lines, i)
            add_para(doc, text)
            i = j
            continue

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                r.font.name = "Consolas"
                r.font.size = Pt(9.5)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F2F2")
                p._p.get_or_add_pPr().append(shd)
            i += 1
            continue
        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("## Mục lục"):
            add_para(doc, "Mục lục", style="Heading 1")
            add_toc_field(doc)
            doc.add_page_break()
            i += 1
            skip_manual_toc = True
            continue
        if skip_manual_toc and re.match(r"^\d+\. \[", line.strip()):
            i += 1
            continue
        if skip_manual_toc and not line.strip():
            i += 1
            continue
        if skip_manual_toc and line.startswith("---"):
            skip_manual_toc = False
            i += 1
            continue

        if line.startswith("#### "):
            add_para(doc, line[5:], style="Heading 3")
            i += 1
            continue
        if line.startswith("### "):
            add_para(doc, line[4:], style="Heading 2")
            i += 1
            continue
        if line.startswith("## "):
            add_para(doc, line[3:], style="Heading 1")
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue  # tiêu đề chính đã xử lý ở trang bìa

        if line.strip() == "---":
            i += 1
            continue

        if line.strip().startswith("!["):
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
            if m:
                alt, src = m.group(1), m.group(2)
                img_path = (ROOT / src).resolve()
                if img_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Cm(15))
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cap.add_run(alt)
                    cr.italic = True
                    cr.font.size = Pt(10)
                    cr.font.color.rgb = GREY
            i += 1
            continue

        if line.strip().startswith("|"):
            table_rows, new_i = parse_table_block(lines, i)
            if table_rows:
                add_table(doc, table_rows[0], table_rows[1:])
            i = new_i
            continue

        if line.strip().startswith("> "):
            add_para(doc, line.strip()[2:], italic=True, color=GREY)
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            text, j = consume_para_lines(lines, i)
            add_para(doc, re.sub(r"^\d+\.\s", "", text), style="List Number")
            i = j
            continue
        if line.strip().startswith("- "):
            text, j = consume_para_lines(lines, i)
            add_para(doc, text[2:] if text.startswith("- ") else text, style="List Bullet")
            i = j
            continue

        if line.strip().startswith("<div") or line.strip().startswith("</div>") or line.strip().startswith("&nbsp;"):
            i += 1
            continue

        if line.strip().startswith("**") and line.strip().endswith("**") and line.count("**") == 2:
            add_para(doc, line.strip(), bold=True)
            i += 1
            continue

        add_para(doc, line.strip())
        i += 1


def build():
    md = MD_PATH.read_text(encoding="utf-8")
    # Bỏ frontmatter YAML
    md = re.sub(r"^---\n.*?\n---\n", "", md, flags=re.DOTALL)
    # Bỏ khối cover <div align="center">...</div> đầu tiên (đã build thủ công)
    md = re.sub(r"<div align=\"center\">\n\n# ỨNG DỤNG.*?</div>", "", md, count=1, flags=re.DOTALL)
    # Bỏ khối kết thúc <div align="center"> cuối file (Hết báo cáo) -> xử lý riêng
    tail_match = re.search(r"<div align=\"center\">\n\n\*— Hết báo cáo.*?</div>", md, flags=re.DOTALL)
    tail_text = tail_match.group(0) if tail_match else ""
    if tail_match:
        md = md[: tail_match.start()]

    doc = Document()
    set_base_styles(doc)
    add_page_number_footer(doc)
    build_cover(doc)
    render_body(doc, md)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— HẾT BÁO CÁO —")
    r.bold = True
    r.font.size = Pt(13)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{AUTHOR} · MSSV {MSSV} · {COURSE} · {DATE}")
    r2.italic = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY

    doc.save(OUT_PATH)
    print(f"Da xuat: {OUT_PATH}")


if __name__ == "__main__":
    build()
